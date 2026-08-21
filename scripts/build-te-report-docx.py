#!/usr/bin/env python3
"""Build the Lak Si intersection operational analysis as a Word document
from the raw survey JSON extracted from หลักสี่(500) 69_05_20.xlsx."""

from __future__ import annotations

import json
import math
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path("/workspace")
SURVEY = json.loads((ROOT / "src/data/survey.json").read_text())
OUT = ROOT / "public/docs/รายงานวิเคราะห์สภาพการจราจรทางแยก_แยกหลักสี่_สัญญาณไฟ500.docx"
ART = ROOT / "artifacts/รายงานวิเคราะห์สภาพการจราจรทางแยก_แยกหลักสี่_สัญญาณไฟ500.docx"

NAVY = RGBColor(0x1F, 0x33, 0x48)
MUTED = RGBColor(0x6B, 0x64, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0x9B, 0x2C, 0x2C)
GREEN = RGBColor(0x2F, 0x5D, 0x43)
FONT = "TH Sarabun New"
FONT_LATIN = "Calibri"

SAT_LEFT = 1710.0


def th(n: float, d: int = 0) -> str:
    return f"{n:,.{d}f}"


def set_run_font(run, size=16, bold=False, color=NAVY, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_LATIN
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:cs"), FONT)


def shade(cell, hex_color: str):
    tc = cell._teCell if hasattr(cell, "_teCell") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C4BBA8")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, bold=False, size=14, center=False, color=NAVY):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    set_cell_border(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=12, color=MUTED)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def signal_delay(cycle, g_over_c, volume, s=SAT_LEFT):
    cap = max(1.0, s * g_over_c)
    x = volume / cap
    xu = min(1.0, max(0.0, x))
    d1 = (0.5 * cycle * (1 - g_over_c) ** 2) / (1 - xu * g_over_c)
    t = 0.25
    d2 = 900 * t * (x - 1 + math.sqrt((x - 1) ** 2 + (8 * 0.5 * x) / (cap * t)))
    return max(0.0, d1 + d2), cap, x


def yield_capacity(vc, tc=6.5, tf=3.3):
    vc = max(1.0, vc)
    a = math.exp((-vc * tc) / 3600)
    b = 1 - math.exp((-vc * tf) / 3600)
    return (vc * a) / max(1e-6, b)


def twsc_delay(volume, capacity, t=0.25):
    c = max(1.0, capacity)
    x = volume / c
    d1 = 3600 / c
    d2 = 900 * t * (x - 1 + math.sqrt((x - 1) ** 2 + (3600 / c) * x / (450 * t)))
    return max(0.0, d1 + d2)


def los_sig(d):
    return "A" if d <= 10 else "B" if d <= 20 else "C" if d <= 35 else "D" if d <= 55 else "E" if d <= 80 else "F"


def los_un(d):
    return "A" if d <= 10 else "B" if d <= 15 else "C" if d <= 25 else "D" if d <= 35 else "E" if d <= 50 else "F"


class Report:
    def __init__(self):
        self.doc = Document()
        self._setup()
        self.s = SURVEY
        self.demand = 342
        self.conflict = 523
        d_geom, cap_geom, x_geom = signal_delay(180, 0.18, self.demand)
        d_sig, cap_sig, x_sig = signal_delay(180, 78 / 180, self.demand)
        free_cap = yield_capacity(self.conflict)
        free_d = twsc_delay(self.demand, free_cap)
        self.geom = (d_geom, cap_geom, x_geom)
        self.sig = (d_sig, cap_sig, x_sig)
        self.free = (free_d, free_cap, self.demand / free_cap)

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(2.0)
        header = sec.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = header.add_run("สำนักการจราจรและขนส่ง  ·  รายงานวิเคราะห์สภาพการจราจรทางแยก  ·  ตู้ 500")
        set_run_font(r, size=12, color=MUTED)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("แยกหลักสี่พลาซ่า  ·  หน้า ")
        set_run_font(r, size=12, color=MUTED)
        add_page_number(footer)
        r = footer.add_run("  ·  ข้อมูลดิบ หลักสี่(500) 69_05_20.xlsx")
        set_run_font(r, size=12, color=MUTED)
        styles = self.doc.styles
        styles["Normal"].font.name = FONT_LATIN
        styles["Normal"].font.size = Pt(16)

    def p(self, text, *, size=16, bold=False, center=False, space_after=8, color=NAVY, italic=False):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color, italic=italic)
        return para

    def h(self, text, level=1):
        size = {1: 20, 2: 18, 3: 16}[level]
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        set_run_font(run, size=size, bold=True, color=NAVY)
        return para

    def table(self, headers, rows, col_widths=None, header_fill="1F3348"):
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        for i, h in enumerate(headers):
            cell_text(tbl.rows[0].cells[i], h, bold=True, size=13, center=True, color=WHITE)
            shade(tbl.rows[0].cells[i], header_fill)
        for r_i, row in enumerate(rows):
            fill = "FBFAf6" if r_i % 2 == 0 else "ECE7DB"
            for c_i, val in enumerate(row):
                cell_text(
                    tbl.rows[r_i + 1].cells[c_i],
                    val,
                    size=13,
                    center=c_i > 0,
                    bold=(str(row[0]).startswith("EB5") if c_i == 0 else False),
                )
                shade(tbl.rows[r_i + 1].cells[c_i], fill)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return tbl

    def caption(self, text):
        self.p(text, size=12, italic=True, color=MUTED, space_after=10)

    def picture(self, path: Path, width_cm=16.5, caption=""):
        if not path.exists():
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        if caption:
            self.caption(caption)

    def cover(self):
        for _ in range(2):
            self.doc.add_paragraph()
        self.p("กรุงเทพมหานคร", size=16, center=True, bold=True, space_after=2)
        self.p("สำนักการจราจรและขนส่ง", size=18, center=True, bold=True, space_after=18)
        self.p("รายงานวิเคราะห์สภาพการจราจรทางแยก", size=26, center=True, bold=True, space_after=4)
        self.p("Intersection Operational Analysis", size=16, center=True, italic=True, space_after=16)
        self.p("แยกหลักสี่พลาซ่า (แยกไอทีสแควร์)", size=20, center=True, bold=True, space_after=4)
        self.p("ถนนแจ้งวัฒนะ ตัด ถนนกำแพงเพชร 6", size=16, center=True, space_after=4)
        self.p("สัญญาณไฟจราจร หมายเลข 500  ·  รหัส N176", size=16, center=True, space_after=20)
        site = self.s["site"]
        memo = self.s["memo"]
        meta = [
            ("วันสำรวจ", f"{site['date']} ({site['dow']})  อากาศ{site['weather']}"),
            ("พิกัด", f"{site['lat']}°N  {site['lng']}°E"),
            ("บันทึกต้นเรื่อง", f"{memo['no']}  ลงวันที่ {memo['date']}"),
            ("จาก / ถึง", f"{memo['from']}  →  {memo['to']}"),
            ("รับเรื่อง สจส.", f"{memo['recvSjs']}  {memo['recvDate']}  ·  {memo['stat']}"),
            ("แหล่งข้อมูลดิบ", "หลักสี่(500) 69_05_20.xlsx  (แผ่น Header, volume, Queue, Report, Analysis)"),
            ("วิธีวิเคราะห์", "HCM 7  ·  Webster  ·  Harders  ·  ITE Recommended Practice"),
        ]
        self.table(["รายการ", "รายละเอียด"], meta)
        self.p(
            "เอกสารนี้จัดทำจากข้อมูลดิบที่สำรวจ 20 พฤษภาคม 2569 ไม่ใช่คำสั่งติดตั้งจนกว่าสำนักการจราจรและขนส่งจะรับรอง "
            "และไม่ทดแทนแบบก่อสร้าง 1:500",
            size=14,
            italic=True,
            color=MUTED,
        )
        self.doc.add_page_break()

    def toc(self):
        self.h("สารบัญ")
        items = [
            "บทที่ 1  บทนำ",
            "บทที่ 2  พื้นที่ศึกษาและสภาพปัจจุบัน",
            "บทที่ 3  วิธีการศึกษาและมาตรฐาน",
            "บทที่ 4  ผลการสำรวจปริมาณจราจร (จากข้อมูลดิบ)",
            "บทที่ 5  แถวคอยและระดับบริการปัจจุบัน",
            "บทที่ 6  วินิจฉัยปัญหา",
            "บทที่ 7  ทางเลือกการปรับปรุง",
            "บทที่ 8  วิเคราะห์ทางเลือกผายปากเลี้ยวซ้าย",
            "บทที่ 9  แบบสัญญาณไฟวงกลมและตารางขัดแย้ง",
            "บทที่ 10  ข้อเสนอและแผนดำเนินงาน",
            "บทที่ 11  ข้อจำกัดและความไม่แน่นอน",
            "บทที่ 12  เอกสารอ้างอิง",
            "ภาคผนวก ก  ตารางปริมาณ 15 นาที ทิศ EB5",
            "ภาคผนวก ข  ตารางปริมาณรายชั่วโมงทุกทิศ",
        ]
        for it in items:
            self.p(it, size=16, space_after=4)
        self.doc.add_page_break()

    def ch1(self):
        self.h("บทที่ 1  บทนำ")
        self.h("1.1  หลักการและเหตุผล", 2)
        self.p(
            "เขตหลักสี่มีหนังสือถึงผู้อำนวยการสำนักการจราจรและขนส่ง ขอคำนวณปริมาณรถบริเวณแยกหลักสี่ "
            "ถนนแจ้งวัฒนะตัดถนนกำแพงเพชร 6 เพื่อประกอบโครงการปาดผายมุมทางเลี้ยวซ้าย ให้มีรัศมีวงเลี้ยวเพิ่มขึ้น "
            "แก้ไขปัญหาจราจรติดขัดบริเวณถนนแจ้งวัฒนะเลี้ยวซ้ายเข้าถนนกำแพงเพชร 6 ด้านทิศเหนือ"
        )
        self.p(
            "รายงานวิเคราะห์ฉบับก่อนที่แนบมาด้วย (รายงานวิเคราะห์สภาพการจราจรทางแยก_ฉบับสมบูรณ์) "
            "ใช้ทิศ NB เลี้ยวซ้าย ซึ่งไม่ตรงบันทึกข้อความ และมีค่า PCU บางช่วงผิดปกติ "
            "รายงานฉบับนี้ใช้ทิศทางตามบันทึก คือ EB5 (แจ้งวัฒนะขาเข้า เลี้ยวซ้ายเข้ากำแพงเพชร 6 ทิศเหนือ) "
            "และคำนวณใหม่จากแผ่น volume / Analysis / Queue ของไฟล์สำรวจดิบ"
        )
        self.h("1.2  วัตถุประสงค์", 2)
        for t in [
            "หาปริมาณ องค์ประกอบรถ และแถวคอยของทิศทางเป้าหมายและทั้งแยก จากข้อมูลดิบ",
            "ประเมินระดับบริการปัจจุบันของขาเลี้ยวซ้ายตาม Highway Capacity Manual",
            "วิเคราะห์ผลของผายปากร่วมกับช่องซ้ายผ่านตลอด เทียบสถานะที่มีสัญญาณไฟ",
            "กำหนดแบบหัวสัญญาณไฟวงกลมและลูกศร พร้อมตารางขัดแย้งตามทิศจริงในแผนที่",
            "เสนอมาตรการที่ลงสนามได้ โดยเคารพเขต Protection Zone A ของรถไฟฟ้าสายสีชมพู",
        ]:
            self.p("•  " + t, space_after=4)
        self.h("1.3  ขอบเขต", 2)
        self.p(
            "ศึกษาสี่แยกสัญญาณไฟตู้ 500 เท่านั้น ประสานกับตู้ 360 (แยกหลักสี่–วิภาวดี ห่างประมาณ 80 เมตร) "
            "ในระดับ OFFSET ไม่จำลองจุลภาคทั้งสายแจ้งวัฒนะ และไม่ประเมินอุบัติเหตุย้อนหลังเพราะไม่มีแฟ้มอุบัติเหตุประกอบบันทึก"
        )

    def ch2(self):
        self.h("บทที่ 2  พื้นที่ศึกษาและสภาพปัจจุบัน")
        self.h("2.1  ที่ตั้งจากไฟล์สำรวจ (แผ่น Header)", 2)
        site = self.s["site"]
        self.table(
            ["รายการในไฟล์ดิบ", "ค่า"],
            [
                ["ชื่อจุดสำรวจ", site["nameTh"].split("(")[0].strip() + " / ตาม Header: หลักสี่"],
                ["หมายเลขสัญญาณไฟจราจร", site["signalId"]],
                ["พิกัด", f"{site['lat']}  {site['lng']}"],
                ["ถนนที่สำรวจ", "แจ้งวัฒนะ - กำแพงเพชร"],
                ["NB road", site["nb"]],
                ["EB road", site["eb"]],
                ["SB road", site["sb"]],
                ["WB road", site["wb"]],
                ["ประเภททางแยก", site["type"]],
                ["วันสำรวจ", f"{site['date']} ({site['dow']})"],
                ["สภาพอากาศ", site["weather"]],
            ],
        )
        self.p(
            "ตู้ 500 คือแยกหลักสี่พลาซ่า / แยกไอทีสแควร์ ไม่ใช่แยกหลักสี่–วิภาวดี (ตู้ 360) "
            "อาคารไอทีสแควร์อยู่มุมตะวันตกเฉียงใต้ซึ่งเป็นมุมที่จะผายปาก "
            "ทางรถไฟสายเหนืออยู่ทิศตะวันออกของกำแพงเพชร 6 วิภาวดีอยู่ทิศตะวันออก"
        )
        self.h("2.2  ทิศทางจริงในแผนที่ (เหนือขึ้นบน ขับชิดซ้าย)", 2)
        self.table(
            ["เข็มทิศ", "ถนน / แลนด์มาร์ก", "ขาเข้าแยก"],
            [
                ["เหนือ", "กำแพงเพชร 6 ทิศเหนือ", "SB"],
                ["ตะวันออก", "แจ้งวัฒนะไปวิภาวดี · ตู้ 360 ~80 ม. · ทางรถไฟ", "WB"],
                ["ใต้", "กำแพงเพชร 6 ทิศใต้ · สายสีชมพู", "NB"],
                ["ตะวันตก", "แจ้งวัฒนะไปปากเกร็ด / ศูนย์ราชการ", "EB ขาเข้า — ทิศเป้าหมาย"],
            ],
        )
        self.p(
            "ทิศเป้าหมายตามบันทึก: รถบนถนนแจ้งวัฒนะขาเข้า (ตะวันตก→ตะวันออก) เลี้ยวซ้ายเข้ากำแพงเพชร 6 ด้านทิศเหนือ "
            "รหัสนับในไฟล์ดิบคือ EB5"
        )
        self.picture(
            ROOT / "attachments/627218008678859434.jpg",
            caption="ภาพที่ 1  บันทึกข้อความต้นเรื่อง กท ๗๘๑๑/๓๙๗๗ (ไฟล์ที่แนบใน Grok)",
        )
        self.picture(
            ROOT / "attachments/S__22405134.jpg",
            caption="ภาพที่ 2  รูปแบบบริเวณที่จะทำการผายปาก ตามแบบเขตหลักสี่ (ไฟล์ที่แนบใน Grok)",
        )
        self.h("2.3  การควบคุมปัจจุบัน", 2)
        self.p(
            "คอนโทรลเลอร์แบบสเตจ ชนิดหลอด LED บริษัท ศรีทองเจริญ จากแผ่นติดตั้งที่สแกนแนบมา "
            "แผนวันรอบ 180 วินาที เริ่ม 06:00 แผนคืนรอบ 135 วินาที เริ่ม 22:00 ใช้สเตจ 1, 2, 3 และ 8 "
            "ช่อง HEAD TYPE, PHASE CONFLICT, เหลือง/แดงทั้งหมด และ OFFSET ว่างทั้งแผ่น "
            "เขียวที่ลงแผ่นรวมเท่ากับรอบพอดี จึงยังไม่มีช่วงเคลียร์ในตัวเลขโปรแกรม"
        )

    def ch3(self):
        self.h("บทที่ 3  วิธีการศึกษาและมาตรฐาน")
        self.p(
            "โครงรายงานยึด ITE Recommended Practice (Transportation Impact Analyses) ปรับเป็นการศึกษาทางแยกเดี่ยว "
            "และใช้วิธีคำนวณตาม Highway Capacity Manual ฉบับที่ 7"
        )
        pcu = self.s["pcu"]
        self.h("3.1  ค่า PCU จากแผ่น Header ของไฟล์ดิบ", 2)
        self.table(
            ["ชนิดรถในไฟล์", "ค่า PCU"],
            [
                ["CAR (รถยนต์)", str(pcu["c"])],
                ["VAN (ตู้, ปิคอัพ)", str(pcu["vn"])],
                ["BUS (เมล์ใหญ่)", str(pcu["b"])],
                ["TRUCK ใน Header = เมล์เล็ก", str(pcu["mb"])],
                ["HEAVY TRUCK (บรรทุก)", str(pcu["tk"])],
                ["TRI-CYCLE (สามล้อ)", str(pcu["tr"])],
            ],
        )
        self.caption("ที่มา: แผ่น Header เซลล์ PCU VALUES ของไฟล์ หลักสี่(500) 69_05_20.xlsx — ไม่ใช่ค่าที่สมมติ")
        self.h("3.2  วิธีคำนวณ", 2)
        self.table(
            ["รายการ", "วิธี"],
            [
                ["นับปริมาณ", "ช่วง 15 นาที 07:00–19:00 จากแผ่น volume / Analysis แยกทิศ NB1–WB16 และชนิดรถ"],
                ["PHF", "V / (4 × V15max) ตาม HCM"],
                ["หน่วงสัญญาณ", f"Webster d1 + HCM overflow d2 · ความอิ่มตัวเลี้ยวซ้าย {th(SAT_LEFT)} PCU/ชม./ช่อง"],
                ["ซ้ายผ่านตลอด (ให้ทาง)", "Harders / HCM unsignalized · tc = 6.5 วินาที · tf = 3.3 วินาที"],
                ["LOS สัญญาณ", "A≤10 B≤20 C≤35 D≤55 E≤80 F เกิน 80 วินาที/คัน"],
                ["LOS ไม่มีไฟ", "A≤10 B≤15 C≤25 D≤35 E≤50 F เกิน 50 วินาที/คัน"],
                ["หัวสัญญาณ", "MUTCD 11 Part 4 สะท้อนเป็นขับซ้าย · UK TSM Chapter 6 · Vienna R.E.2"],
                ["Intergreen", "เหลือง 3 วินาที · แดงทั้งหมด max(2, (W+6)/10)"],
            ],
        )

    def ch4(self):
        self.h("บทที่ 4  ผลการสำรวจปริมาณจราจร")
        self.p(
            "ตัวเลขทั้งหมดในบทนี้รวมและแปลงจากแผ่น Analysis (ผลรวมชนิดรถต่อทิศต่อช่วง 15 นาที) "
            "และแผ่น Report (ปริมาณรายชั่วโมง) ของไฟล์ หลักสี่(500) 69_05_20.xlsx"
        )
        hourly = self.s["hourly"]
        peak = max(hourly, key=lambda h: h["eb5"]["v"])
        ix_peak = max(hourly, key=lambda h: h["tot"])
        am = hourly[0]["eb5"]["v"] + hourly[1]["eb5"]["v"]
        pm = hourly[9]["eb5"]["v"] + hourly[10]["eb5"]["v"]
        q15 = max(self.s["q15"], key=lambda r: r["EB5"])
        self.h("4.1  สรุปทั้งแยกและทิศเป้าหมาย EB5", 2)
        self.table(
            ["ตัวชี้วัด", "ค่าจากข้อมูลดิบ"],
            [
                ["ทั้งแยก 12 ชั่วโมง", f"{th(self.s['ix_12h'])} คัน"],
                ["EB5 12 ชั่วโมง", f"{th(self.s['eb5_12h'])} คัน  ({th(self.s['eb5_12h_pcu'], 1)} PCU)"],
                ["สัดส่วน EB5 ต่อทั้งแยก", f"{self.s['eb5_12h'] / self.s['ix_12h'] * 100:.1f} %"],
                ["เร่งด่วนเช้า 07:00–09:00", f"{th(am)} คัน  (เฉลี่ย {th(am / 2)} คัน/ชม.)"],
                ["เร่งด่วนเย็น 16:00–18:00", f"{th(pm)} คัน  (เฉลี่ย {th(pm / 2)} คัน/ชม.)"],
                ["พีครายชั่วโมง EB5", f"{th(peak['eb5']['v'])} คัน  ช่วง {peak['h'].replace('.', ':')} น.  PHF {peak['eb5']['phf']}"],
                ["พีคทั้งแยก", f"{th(ix_peak['tot'])} คัน  ช่วง {ix_peak['h'].replace('.', ':')} น."],
                ["พีค 15 นาที EB5", f"{th(q15['EB5'])} คัน  ช่วง {q15['t'].replace('.', ':')}  (อัตรา {th(q15['EB5'] * 4)} คัน/ชม.)"],
            ],
        )
        self.p(
            "พีคของทิศ EB5 อยู่ช่วงเที่ยง ไม่ใช่เช้า — การออกแบบด้วยเฉพาะชั่วโมงเร่งด่วนเช้าจะต่ำกว่าความต้องการจริง"
        )
        self.h("4.2  ปริมาณรายขาเข้า 12 ชั่วโมง", 2)
        ap = {"NB": 0, "EB": 0, "SB": 0, "WB": 0}
        for h in hourly:
            for k in ap:
                ap[k] += h["ap"][k]
        labels = {
            "NB": "กำแพงเพชร 6 ทิศเหนือ (รถจากใต้)",
            "EB": "แจ้งวัฒนะ ขาเข้า",
            "SB": "กำแพงเพชร 6 ทิศใต้ (รถจากเหนือ)",
            "WB": "แจ้งวัฒนะ ขาออก (จากแยกหลักสี่)",
        }
        self.table(
            ["ขาเข้า", "ถนน", "คัน", "สัดส่วน"],
            [
                [k, labels[k], th(ap[k]), f"{ap[k] / self.s['ix_12h'] * 100:.1f} %"]
                for k in ("NB", "EB", "SB", "WB")
            ],
        )
        self.h("4.3  ปริมาณรายทิศทาง 12 ชั่วโมง", 2)
        labels_m = {
            "NB1": "NB เลี้ยวซ้าย",
            "NB2": "NB ตรง",
            "NB3": "NB ตรง/ขวา",
            "EB5": "EB เลี้ยวซ้าย (เป้าหมาย)",
            "EB6": "EB ตรง (หลัก)",
            "EB7": "EB ตรง",
            "EB8": "EB เลี้ยวขวา",
            "SB9": "SB เลี้ยวซ้าย",
            "SB10": "SB ตรง",
            "SB11": "SB ตรง/ขวา",
            "WB13": "WB เลี้ยวซ้าย",
            "WB14": "WB ตรง",
        }
        rows = []
        for m, lab in labels_m.items():
            v = sum(h["m"][m]["v"] for h in hourly)
            rows.append([m, lab, th(v)])
        self.table(["รหัสในไฟล์ดิบ", "ทิศทาง", "คัน / 12 ชม."], rows)
        self.caption("ที่มา: รวมจากแผ่น Report / Analysis — แถว EB5 คือทิศตามบันทึกข้อความ")
        self.h("4.4  องค์ประกอบรถทิศ EB5 ทั้งวัน (จากแผ่น Analysis)", 2)
        mix = self.s["mixEb5"]
        tot = mix["c"] + mix["vn"] + mix["b"] + mix["mb"] + mix["tk"] + mix["tr"]
        self.table(
            ["ชนิด", "คัน", "สัดส่วน"],
            [
                ["รถยนต์", th(mix["c"]), f"{mix['c'] / tot * 100:.1f} %"],
                ["ตู้, ปิคอัพ", th(mix["vn"]), f"{mix['vn'] / tot * 100:.1f} %"],
                ["เมล์ใหญ่", th(mix["b"]), f"{mix['b'] / tot * 100:.1f} %"],
                ["เมล์เล็ก", th(mix["mb"]), f"{mix['mb'] / tot * 100:.1f} %"],
                ["บรรทุก", th(mix["tk"]), f"{mix['tk'] / tot * 100:.1f} %"],
                ["สามล้อ", th(mix["tr"]), f"{mix['tr'] / tot * 100:.1f} %"],
                ["รวม", th(tot), "100 %"],
            ],
        )
        self.p(
            f"รถยนต์และตู้รวม {(mix['c'] + mix['vn']) / tot * 100:.1f} % "
            "รถใหญ่ (เมล์+บรรทุก) มี 14 คันต่อ 12 ชั่วโมง — รัศมีที่กว้างขึ้นยังจำเป็นสำหรับรถเหล่านี้"
        )
        self.h("4.5  ตารางชั่วโมงทิศ EB5", 2)
        rows = []
        for h in hourly:
            share = h["eb5"]["v"] / h["ap"]["EB"] * 100
            rows.append(
                [
                    h["h"].replace(".", ":"),
                    th(h["eb5"]["v"]),
                    th(h["eb5"]["p"], 1),
                    str(h["eb5"]["p15"]),
                    str(h["eb5"]["phf"] if h["eb5"]["phf"] is not None else "–"),
                    f"{share:.1f} %",
                    th(h["tot"]),
                ]
            )
        self.table(
            ["ช่วงเวลา", "คัน", "PCU", "พีค 15 นาที", "PHF", "% ขา EB", "ทั้งแยก"],
            rows,
        )

    def ch5(self):
        self.h("บทที่ 5  แถวคอยและระดับบริการปัจจุบัน")
        self.h("5.1  แถวคอยจากแผ่น Queue ของไฟล์ดิบ (เมตร)", 2)
        rows = []
        for k, lab in (
            ("NB", "กำแพงเพชร 6 จากใต้"),
            ("EB", "แจ้งวัฒนะ ขาเข้า"),
            ("SB", "กำแพงเพชร 6 จากเหนือ"),
            ("WB", "แจ้งวัฒนะ จากหลักสี่"),
        ):
            series = self.s["queue"][k]
            mx = max(q["max"] or 0 for q in series)
            when = next(q["h"] for q in series if (q["max"] or 0) == mx)
            rows.append([lab, th(mx), when])
        self.table(["ขา", "ยาวสุด (ม.)", "ช่วง"], rows)
        self.p(
            "ขา EB ยาวที่สุด 200 เมตร ช่วงเช้า — spillback บนแจ้งวัฒนะขาเข้า "
            "ลิงก์ตะวันออกไปตู้ 360 ยาวเพียงประมาณ 80 เมตร คิวจึงปิดปากแยกคู่ได้ง่าย"
        )
        d_geom, cap_geom, x_geom = self.geom
        d_sig, cap_sig, x_sig = self.sig
        self.h("5.2  ความจุและ LOS ของ EB5 ชั่วโมงออกแบบ", 2)
        self.p(
            "ชั่วโมงออกแบบใช้ 12:00–13:00 ตามพีคจริงในไฟล์ดิบ = 342 คัน/ชม. "
            "อัตราพีค 15 นาทีสูงสุด 94 คัน (16:45–17:00) เทียบเท่า 376 คัน/ชม. "
            "เลี้ยวซ้าย EB5 เป็นเลี้ยวใกล้ในระบบขับซ้าย ไปได้ช่วงวงกลม C เขียว (สเตจ 3+8 = 78/180 วินาที) "
            "แต่มุมแคบบังคับความเร็วต่ำ จึงแสดงสองกรณี"
        )
        self.table(
            ["กรณี", "สมมติ", "ความจุ PCU/ชม.", "v/c", "หน่วง วินาที/คัน", "LOS"],
            [
                [
                    "A จำกัดด้วยเรขา",
                    "C = 180 วินาที  g/C = 0.18",
                    th(cap_geom),
                    f"{x_geom:.2f}",
                    th(d_geom),
                    los_sig(d_geom),
                ],
                [
                    "B เขียวตามแผ่นติดตั้ง",
                    "C = 180 วินาที  g/C = 0.43",
                    th(cap_sig),
                    f"{x_sig:.2f}",
                    th(d_sig),
                    los_sig(d_sig),
                ],
            ],
        )
        self.p(
            "กรณี A สอดคล้องคิว EB 200 เมตรในไฟล์ดิบมากกว่ากรณี B จึงใช้กรณี A เป็นฐานเปรียบเทียบทางเลือก"
        )

    def ch6(self):
        self.h("บทที่ 6  วินิจฉัยปัญหา")
        for t in [
            "เรขาคณิตมุมแคบ — เลี้ยวซ้ายต้องชะลอต่ำกว่า 15 กม./ชม. กินเวลาเขียวของช่องใน ดันคิวแจ้งวัฒนะ",
            "ไม่มีช่องเก็บเลี้ยวซ้าย — EB5 ปนช่องตรง EB6/EB7 สัดส่วนราว 19% ของขาเข้าแจ้งวัฒนะ",
            "คิว EB สูงสุด 200 ม. จากแผ่น Queue ยาวสุดทั้งแยก ช่วงเช้า",
            "ลิงก์ 80 ม. ไปตู้ 360 ไม่มี OFFSET ในแผ่นติดตั้ง — เขียว C 78 วินาทีต่อรอบ 180 เทเข้าลิงก์เกินความจุ",
            "หัวสัญญาณบนแผ่นติดตั้งไม่แยกวงกลม/ลูกศร — เสี่ยงเขียวพร้อมกันของเลี้ยวขวาตัด (B×G และ D×A)",
            "Intergreen ไม่ปรากฏในรอบ — เขียวรวมเท่ากับรอบ รถเคลียร์ตัดเขียวเฟสถัดไปได้",
            "ไม่มีเฟสคนเดิน ทั้งที่มีทางม้าลายขาใต้หน้าห้างไอทีสแควร์ในสเก็ตช์",
        ]:
            self.p("•  " + t, space_after=4)

    def ch7(self):
        self.h("บทที่ 7  ทางเลือกการปรับปรุง")
        self.table(
            ["ทางเลือก", "สาระ", "ผลเบื้องต้น"],
            [
                ["0 ไม่ทำอะไร", "คงมุมแคบและรอบ 180 วินาที", "LOS F กรณีเรขา · คิว EB คงยาว"],
                ["1 จูนไฟอย่างเดียว", "ลดรอบ ใส่ intergreen ประสานตู้ 360", "ช่วยทางตรงได้บ้าง เลี้ยวซ้ายยังติดมุม"],
                ["2 ผายปาก + ซ้ายผ่านตลอด", "ตามแบบเขตหลักสี่ + ให้ทางเข้า NB", "ดึง EB5 ออกจากไฟ ลดแรงเสียดทานช่องตรง"],
                [
                    "3 ข้อ 2 + แก้หัวสัญญาณ (แนะนำ)",
                    "ลูกศร B/D · conflict monitor · คนเดิน",
                    "ครบทั้งความจุและความปลอดภัย",
                ],
            ],
        )
        self.h("7.1  ขนาดจากแบบผายปากที่แนบ", 2)
        self.table(
            ["ขนาดในแบบ", "ความหมาย"],
            [
                ["26.15 ม.", "ความยาวแนวปาดตามถนนแจ้งวัฒนะ"],
                ["11.25 ม.", "ระยะปาดด้านเข้ามุม"],
                ["5.14 ม.", "ทางเท้าคงเหลือช่วงกลางมุม"],
                ["9.77 ม.", "แนวโค้งปากทางด้านออก"],
                ["8.77 ม.", "ช่วงเชื่อมเข้าเกาะ/ช่องเร่ง"],
                ["5.49 ม.", "ความกว้างช่วงคอขวดมุม"],
                ["3.29 ม.", "ช่วงเข้าเกาะสามเหลี่ยม"],
                ["3.00 ม.", "ระยะร่น Protection Zone A สายสีชมพู — ห้ามล้ำ"],
            ],
        )

    def ch8(self):
        self.h("บทที่ 8  วิเคราะห์ทางเลือกผายปาก")
        d_geom, cap_geom, x_geom = self.geom
        free_d, free_cap, free_x = self.free
        self.p(
            "หลังผายปาก แยกช่องซ้ายผ่านตลอดออกจากสัญญาณ ใช้ช่องว่างวิกฤตบนกำแพงเพชร 6 ทิศเหนือ "
            f"ปริมาณขัดแย้งใช้ขา NB ชั่วโมง 12:00–13:00 จากไฟล์ดิบ = {th(self.conflict)} คัน/ชม."
        )
        self.table(
            ["ตัวชี้วัด", "ปัจจุบัน (กรณี A)", "หลังผายปาก (ให้ทาง)"],
            [
                ["ความจุ", f"{th(cap_geom)} PCU/ชม.", f"{th(free_cap)} PCU/ชม."],
                ["ความต้องการ", f"{th(self.demand)} คัน/ชม.", f"{th(self.demand)} คัน/ชม."],
                ["v/c", f"{x_geom:.2f}", f"{free_x:.2f}"],
                ["หน่วงควบคุม", f"{th(d_geom)} วินาที/คัน", f"{th(free_d)} วินาที/คัน"],
                ["LOS", los_sig(d_geom), los_un(free_d)],
            ],
        )
        self.p(
            f"หน่วงต่อคันลดประมาณ {th(max(0, d_geom - free_d))} วินาที "
            f"ความจุเหลือเฟือต่อ {th(self.demand)} คัน/ชม. (v/c ≈ {free_x:.2f}) "
            "ประโยชน์หลักต่อทั้งแยกคือดึงราว 15–20% ของขา EB ออกจากสัญญาณ "
            "ให้ช่องตรงรับคิว 200 ม. ได้ดีขึ้น แนวปาด 26 ม. เพียงพอคิวให้ทาง 2–4 คัน"
        )
        self.h("8.1  เงื่อนไขออกแบบประกอบ", 2)
        for t in [
            "ช่องเร่งบนกำแพงเพชร 6 ทิศเหนือ 40–60 ม. หากเขตทางพอ",
            "ป้ายให้ทาง + ลูกศรซ้ายผ่านตลอด + ห้ามจอดปากทาง",
            "ย้ายทางม้าลายพ้นปากโค้ง หรือทำเกาะหลบภัย เพราะซ้ายอิสระตัดคนข้ามมากขึ้น",
            "ตรวจ swept path รถเมล์/บรรทุกจากแบบ 1:500 ก่อนก่อสร้าง (14 คัน/วันในทิศนี้)",
            "ห้ามล้ำ Protection Zone A ระยะร่น 3.00 ม. ตามแบบ",
        ]:
            self.p("•  " + t, space_after=4)

    def ch9(self):
        self.h("บทที่ 9  แบบสัญญาณไฟวงกลมและตารางขัดแย้ง")
        self.p(
            "ผายปากอย่างเดียวไม่แก้ความปลอดภัยของหัวสัญญาณ บทนี้กำหนดไฟวงกลม (Full Green, FG) "
            "และลูกศร (Green Arrow, GA) ตามทิศจริงเหนือขึ้นบน ระบบขับซ้าย "
            "อ้างอิง MUTCD 11 Part 4 (สะท้อนกระจก) และ UK Traffic Signs Manual Chapter 6 ซึ่งตรงกับคอนโทรลเลอร์แบบสเตจ"
        )
        self.h("9.1  ความหมายสากล", 2)
        self.p(
            "ไฟวงกลมเขียว: ไปได้ทุกทิศที่ช่องนั้นอนุญาต ต้องให้ทางคนเดินและรถในแยก "
            "เลี้ยวตัดคู่ตรงข้ามต้องให้ทาง (MUTCD 4A.03)"
        )
        self.p(
            "ไฟลูกศรเขียว: เข้าแยกได้เฉพาะทิศลูกศร ใช้เมื่อต้องการป้องกันการตัดคู่ตรงข้าม "
            "ในระบบขับซ้ายคือเลี้ยวขวา — ที่แยกนี้คือเฟส B (EB เลี้ยวขวาไปใต้) และ D (SB เลี้ยวขวาไปตะวันตก)"
        )
        self.h("9.2  จับคู่เฟสบนแผ่นติดตั้งกับทิศจริงและรหัสนับในไฟล์ดิบ", 2)
        self.table(
            ["เฟส", "ทิศจริง", "รหัสนับ", "หัวที่ต้องใช้", "12 ชม. (คัน)"],
            [
                ["A", "NB ตรง", "NB2", "FG วงกลม", th(sum(h["m"]["NB2"]["v"] for h in self.s["hourly"]))],
                ["B", "EB เลี้ยวขวา", "EB8", "GA ลูกศรขวาเท่านั้น", th(sum(h["m"]["EB8"]["v"] for h in self.s["hourly"]))],
                ["C", "EB ตรง", "EB6+EB7", "FG วงกลมช่องตรง", th(sum(h["m"]["EB6"]["v"] + h["m"]["EB7"]["v"] for h in self.s["hourly"]))],
                ["D", "SB เลี้ยวขวา", "SB11 (ส่วนขวา)", "GA ลูกศรขวาเท่านั้น", th(sum(h["m"]["SB11"]["v"] for h in self.s["hourly"]))],
                ["E", "SB ตรง", "SB10", "FG วงกลมช่องตรง", th(sum(h["m"]["SB10"]["v"] for h in self.s["hourly"]))],
                ["F", "NB เลี้ยวซ้าย (ใกล้)", "NB1", "ไปกับวงกลม A", th(sum(h["m"]["NB1"]["v"] for h in self.s["hourly"]))],
                ["G", "WB ตรง", "WB14", "FG วงกลม", th(sum(h["m"]["WB14"]["v"] for h in self.s["hourly"]))],
            ],
        )
        self.h("9.3  ตาราง PHASE CONFLICT ที่ต้องลงในตู้", 2)
        self.p("Y = ขัดกันห้ามเขียวพร้อมกัน   N = ไปด้วยกันได้   — = เฟสเดียวกัน")
        ids = ["A", "B", "C", "D", "E", "F", "G"]
        mat = {
            "A": dict(A="—", B="Y", C="Y", D="Y", E="N", F="N", G="Y"),
            "B": dict(A="Y", B="—", C="N", D="Y", E="Y", F="Y", G="Y"),
            "C": dict(A="Y", B="N", C="—", D="Y", E="Y", F="N", G="N"),
            "D": dict(A="Y", B="Y", C="Y", D="—", E="N", F="Y", G="Y"),
            "E": dict(A="N", B="Y", C="Y", D="N", E="—", F="N", G="Y"),
            "F": dict(A="N", B="Y", C="N", D="Y", E="N", F="—", G="N"),
            "G": dict(A="Y", B="Y", C="N", D="Y", E="Y", F="N", G="—"),
        }
        rows = [[r] + [mat[r][c] for c in ids] for r in ids]
        self.table(["เฟส"] + ids, rows)
        self.p(
            "จุดตาย: ถ้า C เป็นวงกลมเสาเดียวครอบช่องเลี้ยวขวา ตอนสเตจ 3 (C+G เขียว) รถ EB จะเลี้ยวขวาตัด WB ตรง (B×G) "
            "ถ้า E เป็นวงกลมเสาเดียวครอบช่องเลี้ยวขวา ตอนสเตจ 1 (A+E เขียว) รถ SB จะเลี้ยวขวาตัด NB ตรง (D×A) "
            "แก้ด้วยหัวลูกศรที่ B และ D และแดงลูกศรตอนสเตจที่คู่ตรงข้ามเขียว"
        )
        self.h("9.4  Intergreen และแผนเวลาที่แนะนำ", 2)
        self.table(
            ["จาก → ถึง", "เหลือง", "แดงทั้งหมด", "รวม", "เหตุผล"],
            [
                ["สเตจ 1 → 2", "3", "2", "5", "ตัด A — E ค้างเขียว (overlap)"],
                ["สเตจ 2 → 3", "3", "4", "7", "เคลียร์ขวางแจ้งวัฒนะ กว้าง ~30 ม."],
                ["สเตจ 3 → 8", "3", "4", "7", "ตัด G จากลิงก์ตู้ 360 ก่อนเปิดลูกศร B"],
                ["สเตจ 8 → 1", "3", "3", "6", "C+B ออก  A+E เข้า"],
            ],
        )
        self.table(
            ["สเตจ", "เขียวบนแผ่น วัน/คืน", "เขียวที่แนะนำ วัน/คืน", "เฟส"],
            [
                ["1 เหนือ–ใต้ ตรง", "77 / 65", "54 / 46", "A + E"],
                ["2 SB ตรงค้าง + ขวา D", "25 / 20", "17 / 14", "D + E"],
                ["3 ออก–ตก ตรง", "38 / 20", "26 / 14", "C + G"],
                ["8 EB ตรงค้าง + ขวา B", "40 / 30", "28 / 21", "C + B"],
                ["Intergreen รวม", "0 / 0", "25 / 25", "เหลือง+แดงทั้งหมด"],
                ["รอบ", "180 / 135", "150 / 120", "DAY CODE 10  OFFSET 8–12 วินาที"],
            ],
        )
        self.p(
            "ตาราง FT บนแผ่นเดิมติ๊กขัดทุกคู่สเตจ ซึ่งตัด overlap ให้เหลือช่อง (1,2) และ (3,8) เป็น Non-conflict"
        )

    def ch10(self):
        self.h("บทที่ 10  ข้อเสนอและแผนดำเนินงาน")
        d_geom, _, _ = self.geom
        free_d, _, free_x = self.free
        hourly = self.s["hourly"]
        am = hourly[0]["eb5"]["v"] + hourly[1]["eb5"]["v"]
        pm = hourly[9]["eb5"]["v"] + hourly[10]["eb5"]["v"]
        self.p("เสนอทางเลือก 3: ผายปากตามแบบเขตหลักสี่ + ซ้ายผ่านตลอด EB5 + แก้หัวสัญญาณและจังหวะไฟ")
        self.h("10.1  คำตอบสามข้อในบันทึกข้อความ", 2)
        self.p(
            f"ข้อ 1 ปริมาณช่วงเร่งด่วน — เช้า 07:00–09:00 รวม {th(am)} คัน เย็น 16:00–18:00 รวม {th(pm)} คัน "
            f"พีคจริงของทิศนี้คือ 12:00–13:00 ที่ {th(self.demand)} คัน/ชม. จากแผ่น Report"
        )
        self.p(
            f"ข้อ 2 ใช้เส้นทางนี้ทั้งวัน — {th(self.s['eb5_12h'])} คัน ใน 12 ชั่วโมง "
            f"({th(self.s['eb5_12h_pcu'], 1)} PCU) คิดเป็น {self.s['eb5_12h'] / self.s['ix_12h'] * 100:.1f}% ของรถทั้งแยก"
        )
        self.p(
            f"ข้อ 3 ช่วยได้มากน้อย — ความจุเหลือเฟือ (v/c {free_x:.2f} LOS {los_un(free_d)}) "
            f"หน่วงลดจาก {th(d_geom)} เหลือ {th(free_d)} วินาที/คัน และดึง 15–20% ของขา EB ออกจากไฟ "
            "ซึ่งเป็นขาที่คิวยาวที่สุดในแผ่น Queue (200 ม.)"
        )
        self.h("10.2  ลำดับงาน", 2)
        for i, t in enumerate(
            [
                "รฟม. ยืนยันเขต Protection Zone A ก่อนก่อสร้าง",
                "สำนักการโยธาวัด swept path และช่องเร่ง NB",
                "สจส. ลงโปรแกรมตู้ตามบทที่ 9 (หัว GA ที่ B/D, ตารางขัดแย้ง, intergreen, OFFSET)",
                "เปิดซ้ายผ่านตลอดพร้อมป้ายให้ทาง หลังงานโยธาและหัวสัญญาณเสร็จ",
                "สำรวจหลังเปิด 15 นาที × 1 วันทำการ เพื่อเทียบคิว EB กับฐาน 200 ม. ในไฟล์ดิบ",
            ],
            1,
        ):
            self.p(f"{i}.  {t}", space_after=4)
        self.h("10.3  รายการตรวจก่อนติดตั้งสัญญาณ", 2)
        checks = [
            "หัว B และ D เป็นลูกศร 300 มม. สามดวง ไม่ใช่วงกลม",
            "หัว C และ E เป็นวงกลมเฉพาะช่องตรง มองไม่เห็นจากช่องเลี้ยวขวา",
            "แดงลูกศร B ติดตลอดสเตจ 1, 2, 3 (เขียวเฉพาะสเตจ 8)",
            "แดงลูกศร D ติดตลอดสเตจ 1, 3, 8 (เขียวเฉพาะสเตจ 2)",
            "ตาราง PHASE CONFLICT ลง Y ตามแมทริกซ์บทที่ 9 ก่อนเปิด conflict monitor",
            "ตาราง FT: คู่ (1,2) และ (3,8) เป็น Non-conflict",
            "เหลือง 3 วินาที + แดงทั้งหมดตามตาราง intergreen รวม 25 วินาทีต่อรอบ",
            "รอบกลางวัน 150 วินาที (ไม่ใช่ 180 ที่เขียวเต็มรอบโดยไม่มีเหลือง)",
            "OFFSET จูนตู้ 360 — วัด travel time บนลิงก์ 80 ม. จริง",
            "หัวคนเดินที่ทางม้าลายขาใต้ เดินขนานสเตจ 3 และ 8",
            "มองจากช่องเลี้ยวขวา EB แล้วต้องไม่เห็นวงกลมเขียวของ C ตอนสเตจ 3",
        ]
        for i, t in enumerate(checks, 1):
            self.p(f"{i}.  {t}", space_after=3)

    def ch11(self):
        self.h("บทที่ 11  ข้อจำกัดและความไม่แน่นอน")
        for t in [
            "สำรวจวันพุธวันเดียว (20 พ.ค. 2569) ไม่ครอบคลุมเสาร์–อาทิตย์และวันฝน",
            "ไม่มีแฟ้มอุบัติเหตุ จึงไม่ทำ before-after ด้านความปลอดภัยเชิงสถิติ",
            "ความจุกรณี A ใช้ g/C ประสิทธิผล 0.18 จากเรขา ไม่มีวัด saturation flow ที่สนาม",
            "ปริมาณขัดแย้งของซ้ายผ่านตลอดใช้ขา NB รวมจากไฟล์ดิบ ไม่แยกช่องรับจริง",
            "ไม่จำลองจุลภาคผลต่อตู้ 360 — OFFSET ที่เสนอต้องวัด travel time จริง",
            "ระยะเคลียร์ intergreen ประมาณจากความกว้างแยกทั่วไป ต้องวัด W จริงก่อนลงตู้",
            "รายงานฉบับก่อนที่แนบมาด้วยไม่ถูกใช้เป็นแหล่งปริมาณ เพราะทิศและ PCU ไม่ตรงไฟล์ดิบ",
        ]:
            self.p("•  " + t, space_after=4)

    def ch12(self):
        self.h("บทที่ 12  เอกสารอ้างอิง")
        refs = [
            "TRB. Highway Capacity Manual, 7th Edition.",
            "ITE. Transportation Impact Analyses for Site Development. Recommended Practice.",
            "FHWA. Traffic Signal Timing Manual. Chapters 4–6.",
            "FHWA. Manual on Uniform Traffic Control Devices, 11th Edition. Part 4.",
            "UK DfT. Traffic Signs Manual Chapter 6 — Traffic Control.",
            "UNECE. Vienna Convention on Road Signs and Signals, R.E.2.",
            "Webster, F.V. Traffic Signal Settings. Road Research Technical Paper No. 39.",
            "Harders, J. Die Leistungsfähigkeit nicht signalgeregelter städtischer Verkehrsknoten.",
            "สำนักการจราจรและขนส่ง. ไฟล์สำรวจ หลักสี่(500) 69_05_20.xlsx, 20 พฤษภาคม 2569. แผ่น Header, volume, Queue, Report, Analysis.",
            f"เขตหลักสี่. บันทึกข้อความ {self.s['memo']['no']} ลงวันที่ {self.s['memo']['date']}.",
            "บริษัท ศรีทองเจริญ จำกัด. แผ่นติดตั้งเครื่องควบคุมสัญญาณไฟจราจร ตู้ 500 (ไฟล์สแกนที่แนบ).",
            "เขตหลักสี่. รูปแบบบริเวณที่จะทำการผายปาก แยกหลักสี่ (แบบขอใช้พื้นที่ ที่แนบ).",
        ]
        for i, t in enumerate(refs, 1):
            self.p(f"{i}.  {t}", size=14, space_after=4)

    def appendix_a(self):
        self.doc.add_page_break()
        self.h("ภาคผนวก ก  ตารางปริมาณ 15 นาที ทิศ EB5")
        self.p("ที่มา: แผ่น Analysis ของไฟล์ หลักสี่(500) 69_05_20.xlsx คอลัมน์ EB5 รวมทุกชนิดรถ")
        rows = []
        for r in self.s["q15"]:
            rows.append([r["t"].replace(".", ":"), th(r["EB5"]), th(r["eb5p"], 1), th(r["tot"])])
        self.table(["ช่วงเวลา", "EB5 คัน", "EB5 PCU", "ทั้งแยก คัน"], rows)

    def appendix_b(self):
        self.h("ภาคผนวก ข  ปริมาณรายชั่วโมงทุกทิศจากแผ่น Report")
        self.p("หน่วย: คัน/ชม. ไม่รวมจักรยานยนต์ ตามหัวตารางในไฟล์ดิบ")
        keys = ["NB1", "NB2", "NB3", "EB5", "EB6", "EB7", "EB8", "SB9", "SB10", "SB11", "WB13", "WB14"]
        header = ["ช่วง"] + keys
        rows = []
        for h in self.s["hourly"]:
            rows.append([h["h"].replace(".", ":")] + [th(h["m"][k]["v"]) for k in keys])
        tot = ["รวม 12 ชม."] + [th(sum(h["m"][k]["v"] for h in self.s["hourly"])) for k in keys]
        rows.append(tot)
        self.table(header, rows)

    def build(self):
        self.cover()
        self.toc()
        self.ch1()
        self.ch2()
        self.ch3()
        self.ch4()
        self.ch5()
        self.ch6()
        self.ch7()
        self.ch8()
        self.ch9()
        self.ch10()
        self.ch11()
        self.ch12()
        self.appendix_a()
        self.appendix_b()
        OUT.parent.mkdir(parents=True, exist_ok=True)
        ART.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(OUT)
        self.doc.save(ART)
        print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")
        print(f"wrote {ART} ({ART.stat().st_size} bytes)")


if __name__ == "__main__":
    Report().build()
